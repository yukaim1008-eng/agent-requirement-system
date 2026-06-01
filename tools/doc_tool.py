"""
PRD 导出工具 —— 把 Markdown 格式的 PRD 文本导出为 Word .docx，存到 output/

设计选择：导出在 crew 跑完后由代码确定性完成，不交给 LLM 当工具调用
（让 LLM 处理整篇文档的文件 IO 不可靠，确定性导出更稳）
支持：标题(# ## ###)、有序/无序列表、Markdown 表格、行内 ** 加粗
"""
import os
import re
from datetime import date

from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")


def _strip_md(text: str) -> str:
    """去掉行内 markdown 标记（** 加粗、` 代码），保留纯文本。"""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def export_prd_to_docx(markdown_text: str, filename: str | None = None) -> str:
    """
    把 Markdown 格式的 PRD 写入 output/<filename>.docx，返回文件绝对路径。
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    if not filename:
        filename = f"PRD_{date.today().isoformat()}.docx"
    if not filename.endswith(".docx"):
        filename += ".docx"
    path = os.path.join(OUTPUT_DIR, filename)

    doc = Document()
    table_buffer = []

    def flush_table():
        """把累积的 Markdown 表格行渲染成 Word 表格。"""
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = [[c.strip() for c in r.strip().strip("|").split("|")] for r in table_buffer]
        # 第二行若是 |:---|---| 分隔符，则去掉
        if len(rows) >= 2 and all(set(c) <= set(":- ") and c for c in rows[1]):
            header, body = rows[0], rows[2:]
        else:
            header, body = rows[0], rows[1:]
        table = doc.add_table(rows=1, cols=len(header))
        table.style = "Table Grid"
        for j, h in enumerate(header):
            table.rows[0].cells[j].text = _strip_md(h)
        for r in body:
            cells = table.add_row().cells
            for j, c in enumerate(r):
                if j < len(cells):
                    cells[j].text = _strip_md(c)
        doc.add_paragraph("")
        table_buffer = []

    for line in markdown_text.split("\n"):
        stripped = line.strip()
        # 表格行先缓存
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            continue
        flush_table()  # 非表格行，先把缓存的表格输出

        if not stripped or stripped == "---":
            continue
        if stripped.startswith("# "):
            doc.add_heading(_strip_md(stripped[2:]), level=0)
        elif stripped.startswith("## "):
            doc.add_heading(_strip_md(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(_strip_md(stripped[4:]), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(_strip_md(stripped[5:]), level=3)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(_strip_md(stripped[2:]), style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(_strip_md(re.sub(r"^\d+\.\s", "", stripped)), style="List Number")
        else:
            doc.add_paragraph(_strip_md(stripped))

    flush_table()  # 收尾：文档末尾若是表格
    doc.save(path)
    return path
